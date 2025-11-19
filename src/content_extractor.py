"""
Document and Web Content Extractor
Handles PDF, DOCX, and web page scraping for job descriptions
"""
import re
from typing import Optional
import requests
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO


class ContentExtractor:
    """Extract text from various sources"""
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_bytes: PDF file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            pdf_file = BytesIO(file_bytes)
            pdf_reader = PdfReader(pdf_file)
            
            text_content = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
            
            return "\n\n".join(text_content)
        
        except Exception as e:
            raise ValueError(f"Failed to extract PDF content: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_bytes: bytes) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_bytes: DOCX file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            docx_file = BytesIO(file_bytes)
            doc = Document(docx_file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_content.append(row_text)
            
            return "\n\n".join(text_content)
        
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX content: {str(e)}")
    
    @staticmethod
    def extract_from_url(url: str) -> str:
        """
        Scrape text content from a web page
        
        Args:
            url: Web page URL
            
        Returns:
            Extracted text content
        """
        try:
            # Add headers to mimic browser
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "header", "footer"]):
                script.decompose()
            
            # Try to find main content area (common job posting elements)
            main_content = None
            
            # Try common job posting containers
            selectors = [
                'div[class*="job-description"]',
                'div[class*="job-detail"]',
                'div[class*="posting"]',
                'article',
                'main',
                'div[role="main"]',
                'div[id*="job"]',
                'div[class*="content"]'
            ]
            
            for selector in selectors:
                element = soup.select_one(selector)
                if element:
                    main_content = element
                    break
            
            # If no specific container found, use body
            if not main_content:
                main_content = soup.body
            
            if not main_content:
                raise ValueError("Could not find content in the page")
            
            # Extract text
            text = main_content.get_text(separator='\n', strip=True)
            
            # Clean up multiple newlines and spaces
            text = re.sub(r'\n\s*\n', '\n\n', text)
            text = re.sub(r' +', ' ', text)
            
            # Remove very short lines (likely navigation/footer)
            lines = [line.strip() for line in text.split('\n') if len(line.strip()) > 20]
            cleaned_text = '\n\n'.join(lines)
            
            if not cleaned_text or len(cleaned_text) < 100:
                raise ValueError("Extracted content is too short, likely not a job posting")
            
            return cleaned_text
        
        except requests.RequestException as e:
            raise ValueError(f"Failed to fetch URL: {str(e)}")
        except Exception as e:
            raise ValueError(f"Failed to extract web content: {str(e)}")
    
    @staticmethod
    def detect_file_type(filename: str) -> Optional[str]:
        """
        Detect file type from filename
        
        Args:
            filename: Name of the file
            
        Returns:
            File type: 'pdf', 'docx', or None
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return 'pdf'
        elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
            return 'docx'
        
        return None
    
    @staticmethod
    def is_valid_url(url: str) -> bool:
        """
        Check if string is a valid URL
        
        Args:
            url: String to check
            
        Returns:
            True if valid URL, False otherwise
        """
        url_pattern = re.compile(
            r'^https?://'  # http:// or https://
            r'(?:(?:[A-Z0-9](?:[A-Z0-9-]{0,61}[A-Z0-9])?\.)+[A-Z]{2,6}\.?|'  # domain...
            r'localhost|'  # localhost...
            r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})'  # ...or ip
            r'(?::\d+)?'  # optional port
            r'(?:/?|[/?]\S+)$', re.IGNORECASE
        )
        return url_pattern.match(url) is not None
